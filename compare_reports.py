from docx import Document
import os

rdir = r"C:\AI_precheck\reports"
files = sorted(os.listdir(rdir))

for f in files:
    doc = Document(os.path.join(rdir, f))
    print(f"FILE: {f}")
    for p in doc.paragraphs:
        t = p.text.strip()
        if "Score" in t or "Checked" in t:
            print(f"  {t}")
    
    for tb in doc.tables:
        for row in tb.rows[1:]:
            cells = [x.text.strip() for x in row.cells]
            if len(cells) >= 4:
                status = cells[2] if len(cells) > 2 else ""
                if "Fail" in status or "Warning" in status:
                    item = cells[1][:35]
                    found = cells[3][:40] if len(cells) > 3 else ""
                    print(f"  {status[:10]:10s} | {item:35s} | {found}")
    print("---")
